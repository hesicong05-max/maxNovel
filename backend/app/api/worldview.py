"""Worldview upload and parsing API."""

import io
import logging
import os
import tempfile
import zipfile
from types import SimpleNamespace
from typing import Annotated

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings as app_settings
from app.core.auth import User, get_current_user, get_project_for_owner
from app.core.maintenance import (
    ensure_project_writes_available,
    require_project_writes_available,
)
from app.core.lore_migration_preview import migration_preview_source_checksum
from app.core.legacy_json import read_legacy_object_list
from app.core.project_files import save_worldview_file
from app.core.worldview_parser import worldview_parser
from app.database import get_db
from app.models.lore import ProjectLoreMigrationOperation
from app.models.project import Project, ProjectStatus, Worldview
from app.schemas.models import (
    WorldviewImportRequest,
    WorldviewImportResponse,
    WorldviewResponse,
    WorldviewSaveRequest,
)

router = APIRouter(prefix="/api/worldview", tags=["worldview"])


def _response_object_list(value, field: str) -> list[dict]:
    """Read native or historical JSON-text collections without rewriting them."""
    result = read_legacy_object_list(value)
    if not result.valid:
        raise HTTPException(
            status_code=422,
            detail={
                "code": "WORLDVIEW_LEGACY_JSON_INVALID",
                "message": f"旧世界观字段 {field} 无法安全读取，请先修复原始数据。",
                "retryable": False,
            },
        )
    return result.items

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".txt", ".md", ".markdown", ".doc", ".docx"}
MAX_UPLOAD_SIZE = app_settings.MAX_UPLOAD_SIZE
MAX_DOCX_ENTRIES = 5000
MAX_DOCX_UNCOMPRESSED_SIZE = 50 * 1024 * 1024
MAX_EXTRACTED_TEXT_CHARS = 200_000


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a bounded .docx archive using python-docx."""
    from docx import Document

    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
            entries = archive.infolist()
            uncompressed_size = sum(entry.file_size for entry in entries)
    except zipfile.BadZipFile as exc:
        raise ValueError("文件不是有效的 DOCX 文档") from exc

    if len(entries) > MAX_DOCX_ENTRIES:
        raise ValueError("DOCX 内部文件数量异常")
    if uncompressed_size > MAX_DOCX_UNCOMPRESSED_SIZE:
        raise ValueError("DOCX 解压后内容过大")

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


async def _extract_text_from_doc(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from a .doc file.
    Uses macOS textutil to convert .doc to plain text.
    Falls back to raising an error if textutil is unavailable.
    """
    import asyncio as aio

    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    output_path = tmp_path + ".txt"
    try:
        proc = await aio.create_subprocess_exec(
            "textutil",
            "-convert",
            "txt",
            "-output",
            output_path,
            tmp_path,
            stdout=aio.subprocess.PIPE,
            stderr=aio.subprocess.PIPE,
        )
        try:
            await aio.wait_for(proc.wait(), timeout=30)
        except aio.TimeoutError:
            proc.kill()
            await proc.wait()
            raise HTTPException(
                status_code=422, detail="解析 .doc 文件超时，请尝试转换为 .docx 格式"
            )

        if proc.returncode == 0 and os.path.exists(output_path):
            with open(output_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read().strip()
                if text:
                    return text

        raise HTTPException(
            status_code=422,
            detail="无法解析 .doc 文件，请尝试转换为 .docx 格式后重新上传",
        )
    finally:
        for path in [tmp_path, output_path]:
            if os.path.exists(path):
                os.unlink(path)


@router.post("/{project_id}/upload-file")
async def upload_worldview_file(
    project_id: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
    file: UploadFile = File(...),
):
    """
    Upload a document file (.txt, .md, .doc, .docx) and extract its text content.
    Returns the extracted text for the frontend to review before calling /import.
    """
    # Verify project ownership
    await get_project_for_owner(project_id, current_user, db)

    # Validate file extension
    filename = file.filename or ""
    _, ext = os.path.splitext(filename.lower())
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {ext}，请上传 .txt / .md / .doc / .docx 文件",
        )

    # Read file content in chunks with size limit (prevents OOM on large uploads)
    file_buffer = bytearray()
    while chunk := await file.read(1024 * 1024):  # 1MB chunks
        file_buffer.extend(chunk)
        if len(file_buffer) > MAX_UPLOAD_SIZE:
            max_mb = MAX_UPLOAD_SIZE / (1024 * 1024)
            raise HTTPException(
                status_code=413,
                detail=f"文件过大，最大允许 {max_mb:.0f}MB",
            )
    file_bytes = bytes(file_buffer)
    file_size = len(file_bytes)
    logger.info("File upload: %s (%d bytes)", filename, file_size)

    if file_size < 10:
        raise HTTPException(status_code=400, detail="文件内容为空或过短")

    # Extract text based on file type
    if ext in (".txt", ".md", ".markdown"):
        # Try UTF-8 first, fall back to other encodings
        for encoding in ("utf-8", "gbk", "gb2312", "big5", "latin-1"):
            try:
                text = file_bytes.decode(encoding)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        else:
            text = file_bytes.decode("utf-8", errors="replace")
    elif ext == ".docx":
        try:
            text = _extract_text_from_docx(file_bytes)
        except Exception as e:
            logger.exception("DOCX parsing failed for %s: %s", filename, e)
            raise HTTPException(
                status_code=422,
                detail="解析 .docx 文件失败，请确认文件有效且内容大小正常",
            )
    elif ext == ".doc":
        try:
            text = await _extract_text_from_doc(file_bytes, filename)
        except HTTPException:
            raise
        except Exception as e:
            logger.exception("DOC parsing failed for %s: %s", filename, e)
            raise HTTPException(
                status_code=422,
                detail="解析 .doc 文件失败，请尝试转换为 .docx 格式后重新上传",
            )
    else:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    if not text or len(text.strip()) < 10:
        raise HTTPException(
            status_code=400,
            detail="从文件中提取的文本内容过短，至少需要 10 个字符",
        )
    if len(text) > MAX_EXTRACTED_TEXT_CHARS:
        raise HTTPException(
            status_code=413,
            detail=f"提取文本过长，最多允许 {MAX_EXTRACTED_TEXT_CHARS} 个字符，请精简后重试",
        )

    return {"text": text, "filename": filename, "char_count": len(text)}


@router.post("/{project_id}/import", response_model=WorldviewImportResponse)
async def import_worldview(
    project_id: str,
    data: WorldviewImportRequest,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
):
    """
    Import worldview from a free-form document.
    Uses LLM to extract structured worldview elements from the text.
    Does NOT save to database — returns extracted elements for user review.
    The user then saves via POST /{project_id} with the extracted data.
    """
    project = await get_project_for_owner(project_id, current_user, db)

    if not data.document_text or len(data.document_text.strip()) < 10:
        raise HTTPException(status_code=400, detail="文档内容过短，至少需要 10 个字符")

    try:
        extracted = await worldview_parser.parse_document(
            data.document_text, project.genre.value
        )
    except ValueError as e:
        raise HTTPException(status_code=422, detail=f"解析失败: {e!s}")

    # Count total elements
    total = (
        len(extracted.get("characters", []))
        + len(extracted.get("geography", []))
        + len(extracted.get("factions", []))
        + len(extracted.get("power_system", []))
        + len(extracted.get("history", []))
        + len(extracted.get("conflicts", []))
        + len(extracted.get("special_settings", []))
    )

    return WorldviewImportResponse(
        characters=extracted.get("characters", []),
        geography=extracted.get("geography", []),
        factions=extracted.get("factions", []),
        power_system=extracted.get("power_system", []),
        history=extracted.get("history", []),
        conflicts=extracted.get("conflicts", []),
        special_settings=extracted.get("special_settings", []),
        raw_text=data.document_text,
        source="imported",
        element_count=total,
    )


@router.post("/{project_id}", response_model=WorldviewResponse)
async def set_worldview(
    project_id: str,
    data: WorldviewSaveRequest,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
    _write_gate: Annotated[None, Depends(require_project_writes_available)],
):
    ensure_project_writes_available()

    # Migration and legacy saves share the same lock order.  This prevents an
    # already-started save from replacing the retained source after a migration
    # has materialized or switched the project to relational storage.
    project = await db.scalar(
        select(Project).where(Project.id == project_id).with_for_update()
    )
    if project is None:
        raise HTTPException(status_code=404, detail="项目不存在")
    if project.owner_id is None or project.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="无权操作此项目")
    protected_migration = await db.scalar(
        select(ProjectLoreMigrationOperation).where(
            ProjectLoreMigrationOperation.project_id == project_id,
            ProjectLoreMigrationOperation.status.in_(("validating", "ready")),
        )
    )
    if project.lore_storage_mode == "migrating" or protected_migration is not None:
        raise HTTPException(
            status_code=409,
            detail={
                "code": "WORLDVIEW_SOURCE_READ_ONLY",
                "message": "此项目的旧世界观已进入升级流程，仅可作为历史来源查看。",
                "retryable": False,
            },
        )

    existing_wv = await db.scalar(
        select(Worldview)
        .where(Worldview.project_id == project_id)
        .with_for_update()
    )

    # Parse worldview into elements. The optimistic token is transport metadata
    # and must never be persisted into the author's source document.
    worldview_dict = data.model_dump(exclude={"expected_source_checksum"})
    elements = worldview_parser.parse(worldview_dict)

    checksum_candidate = SimpleNamespace(
        **worldview_dict,
        parsed_elements=elements,
    )
    target_source_checksum = migration_preview_source_checksum(checksum_candidate)

    if existing_wv is not None:
        current_source_checksum = migration_preview_source_checksum(existing_wv)
        if data.expected_source_checksum != current_source_checksum:
            # A lost response may cause the exact saved payload to be retried.
            # Treat that as an idempotent replay, but reject every stale change.
            if target_source_checksum == current_source_checksum:
                save_worldview_file(project_id, existing_wv)
                return WorldviewResponse(
                    id=existing_wv.id,
                    project_id=existing_wv.project_id,
                    **worldview_dict,
                    parsed_elements=elements,
                    source_checksum=current_source_checksum,
                    created_at=existing_wv.created_at,
                )
            raise HTTPException(
                status_code=409,
                detail={
                    "code": "WORLDVIEW_SOURCE_STALE",
                    "message": "服务器上的世界观已发生变化，请先重新加载并核对本地草稿。",
                    "retryable": False,
                    "reload_required": True,
                },
            )

    if existing_wv is None:
        worldview = Worldview(project_id=project_id)
        db.add(worldview)
    else:
        worldview = existing_wv
    worldview.characters = worldview_dict["characters"]
    worldview.geography = worldview_dict["geography"]
    worldview.factions = worldview_dict["factions"]
    worldview.power_system = worldview_dict["power_system"]
    worldview.history = worldview_dict["history"]
    worldview.conflicts = worldview_dict["conflicts"]
    worldview.special_settings = worldview_dict["special_settings"]
    worldview.raw_text = data.raw_text
    worldview.source = data.source
    worldview.parsed_elements = elements

    project.status = ProjectStatus.WORLDVIEW_SET
    ensure_project_writes_available()
    await db.commit()
    await db.refresh(worldview)

    # Persist as independent document file (DB + file dual write)
    save_worldview_file(project_id, worldview)

    return WorldviewResponse(
        id=worldview.id,
        project_id=worldview.project_id,
        characters=data.characters,
        geography=data.geography,
        factions=data.factions,
        power_system=data.power_system,
        history=data.history,
        conflicts=data.conflicts,
        special_settings=data.special_settings,
        raw_text=data.raw_text,
        source=worldview.source,
        parsed_elements=elements,
        source_checksum=migration_preview_source_checksum(worldview),
        created_at=worldview.created_at,
    )


@router.get("/{project_id}", response_model=WorldviewResponse)
async def get_worldview(
    project_id: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
):
    await get_project_for_owner(project_id, current_user, db)

    result = await db.execute(
        select(Worldview).where(Worldview.project_id == project_id)
    )
    worldview = result.scalar_one_or_none()
    if not worldview:
        raise HTTPException(status_code=404, detail="世界观不存在，请先上传")

    return WorldviewResponse(
        id=worldview.id,
        project_id=worldview.project_id,
        characters=_response_object_list(worldview.characters, "characters"),
        geography=_response_object_list(worldview.geography, "geography"),
        factions=_response_object_list(worldview.factions, "factions"),
        power_system=_response_object_list(worldview.power_system, "power_system"),
        history=_response_object_list(worldview.history, "history"),
        conflicts=_response_object_list(worldview.conflicts, "conflicts"),
        special_settings=_response_object_list(worldview.special_settings, "special_settings"),
        raw_text=worldview.raw_text,
        source=worldview.source or "manual",
        parsed_elements=_response_object_list(worldview.parsed_elements, "parsed_elements"),
        source_checksum=migration_preview_source_checksum(worldview),
        created_at=worldview.created_at,
    )


@router.get("/{project_id}/summary")
async def get_worldview_summary(
    project_id: str,
    db: Annotated[AsyncSession, Depends(get_db)],
    current_user: Annotated[User, Depends(get_current_user)],
):
    await get_project_for_owner(project_id, current_user, db)

    result = await db.execute(
        select(Worldview).where(Worldview.project_id == project_id)
    )
    worldview = result.scalar_one_or_none()
    if not worldview:
        raise HTTPException(status_code=404, detail="世界观不存在")

    summary = worldview_parser.summary(
        worldview_parser.normalize_elements(
            _response_object_list(worldview.parsed_elements, "parsed_elements")
        )
    )
    return summary
